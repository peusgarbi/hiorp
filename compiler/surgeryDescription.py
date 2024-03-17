from looper.cirurgiasLooper import Cirurgias
from datetime import datetime
from consts import surgeons
from docx import Document
import os

def preencher_modelo(patient_name: str, surgeon: str, surgery_date: str):
    modelo = Document("templates/DESCRIÇÕES R1/DRA CLAUDIA/A+A.docx")
    data_datetime = datetime.strptime(surgery_date, '%d/%m/%y')
    data_formatada_sem_barras = data_datetime.strftime('%d%m%Y')
    data_formatada = data_datetime.strftime('%d/%m/%Y')

    dados = {
        "<PACIENTE>": patient_name,
        "<DATA>": data_formatada,
    }

    documento = modelo
    for paragrafo in documento.paragraphs:
        for chave, valor in dados.items():
            if chave in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(chave, valor)
                
    for seção in documento.sections:
        for cabecalho in seção.header.tables:
            for row in cabecalho.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for chave, valor in dados.items():
                            if chave in paragraph.text:
                                paragraph.text = paragraph.text.replace(chave, valor)
    
    for tabela in documento.tables:
        for row in tabela.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for chave, valor in dados.items():
                        if chave in paragraph.text:
                            paragraph.text = paragraph.text.replace(chave, valor)

    if not os.path.exists("generated"):
        os.makedirs("generated")
    documento.save(f"generated/{patient_name}     {data_formatada_sem_barras}.docx")

def make_descriptions(cirurgias: Cirurgias, surgeries_date: str) -> None:
    for cirurgia in cirurgias.sala1:
        if cirurgia.cirurgiao in surgeons.we_do_all_surgeons or cirurgia.cirurgiao in surgeons.we_only_do_descriptions_surgeons:
            preencher_modelo(cirurgia.paciente, cirurgia.cirurgiao, surgeries_date)

    for cirurgia in cirurgias.sala2:
        if cirurgia.cirurgiao in surgeons.we_do_all_surgeons or cirurgia.cirurgiao in surgeons.we_only_do_descriptions_surgeons:
            preencher_modelo(cirurgia.paciente, cirurgia.cirurgiao, surgeries_date)

    for cirurgia in cirurgias.sala3:
        if cirurgia.cirurgiao in surgeons.we_do_all_surgeons or cirurgia.cirurgiao in surgeons.we_only_do_descriptions_surgeons:
            preencher_modelo(cirurgia.paciente, cirurgia.cirurgiao, surgeries_date)

    for cirurgia in cirurgias.sala4:
        if cirurgia.cirurgiao in surgeons.we_do_all_surgeons or cirurgia.cirurgiao in surgeons.we_only_do_descriptions_surgeons:
            preencher_modelo(cirurgia.paciente, cirurgia.cirurgiao, surgeries_date)

    for cirurgia in cirurgias.sala5:
        if cirurgia.cirurgiao in surgeons.we_do_all_surgeons or cirurgia.cirurgiao in surgeons.we_only_do_descriptions_surgeons:
            preencher_modelo(cirurgia.paciente, cirurgia.cirurgiao, surgeries_date)
