from docx import Document

def preencher_modelo(paciente: str, data: str):
    modelo = Document('templates/DESCRIÇÕES R1/DRA CLAUDIA/A+A.docx')
    documento = modelo
    dados = {
        '<PACIENTE>': paciente,
        '<DATA>': data,
    }

    for secao in documento.sections:
        # Verifica se o cabeçalho possui tabelas
        if secao.header is not None:
            for tabela in secao.header.tables:
                # Percorre as células da tabela
                for row in tabela.rows:
                    for cell in row.cells:
                        for chave, valor in dados.items():
                            if chave in cell.text:
                                cell.text = cell.text.replace(chave, valor)

    for paragrafo in documento.paragraphs:
        for chave, valor in dados.items():
            if chave in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(chave, valor)

    for tabela in documento.tables:
        for row in tabela.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for chave, valor in dados.items():
                        if chave in paragraph.text:
                            paragraph.text = paragraph.text.replace(chave, valor)
    documento.save("test.docx")
